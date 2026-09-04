from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'relatorio'
OUT.mkdir(exist_ok=True)


def set_cell_shading(cell, fill='EDEDED'):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    cant_split.set(qn('w:val'), 'true')
    tr_pr.append(cant_split)


def set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            set_cell_margins(row.cells[idx])


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])


def add_table(doc, caption, headers, rows, widths=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(caption)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    set_row_cant_split(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(value)
        set_cell_shading(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0, 0, 0)
    for row in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cells[i].paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0, 0, 0)
    if widths:
        set_col_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(11)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run('Nota metodológica. ')
    run.bold = True
    p.add_run(text)
    for r in p.runs:
        r.font.name = 'Arial'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.name = 'Arial'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    return p


def setup(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 1.15
    for name, size in (('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 11)):
        style = doc.styles[name]
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Visagio — Sprint 4 | ')
    add_page_field(footer)
    for r in footer.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0, 0, 0)


def cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('PONTIFÍCIA UNIVERSIDADE CATÓLICA DO RIO DE JANEIRO\n').bold = True
    p.add_run('Projeto de Engenharia de Produção\n\n')
    p.add_run('SIMULAÇÃO DE CENÁRIOS DE MALHA LOGÍSTICA COM IMPACTO DA REFORMA TRIBUTÁRIA\n').bold = True
    p.add_run('Um estudo de caso em cadeias de varejo e indústria')
    for r in p.runs:
        r.font.name = 'Arial'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        r.font.size = Pt(14 if r.bold else 12)
        r.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Hugo Hickman — 230205\nGiovanna Sabatini — 2011088\n\nOrientador: Prof. Igor Tona Peres')
    for r in p.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 0, 0)
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph('Rio de Janeiro\nJunho de 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0, 0, 0)
    doc.add_page_break()


def main():
    doc = Document()
    setup(doc)
    cover(doc)

    heading(doc, 'Resumo executivo', 1)
    add_body(
        doc,
        'Este trabalho avalia cenários discretos de malha logística para duas empresas, considerando custos físicos, tributação parametrizada, estoque, distribuição, transferência e frete. A solução implementada compara uma lista finita de configurações de centros de distribuição (CDs), calcula indicadores de qualidade, risco e robustez e oferece uma camada Monte Carlo exploratória para representar incerteza parametrizada.',
    )
    add_body(
        doc,
        'A principal correção desta versão é separar três objetos que estavam confundidos: o baseline operacional atual, o baseline comparável no mesmo regime tributário do cenário e o resultado simulado. Os savings publicados passam a usar o baseline comparável da própria empresa. O Monte Carlo é descrito como camada probabilística complementar; seus spreads são premissas manuais, não calibração estatística histórica.',
    )
    add_body(
        doc,
        'Na execução canônica, a Empresa 1 tem como melhor cenário dentro do espaço avaliado o candidato `empresa1_candidate_015`, com 14 CDs, custo total de R$ 155.296.806,21 e saving de 1,001985% contra o baseline comparável de reforma de R$ 156.868.606,74. O candidato permanece em todos os quatro perfis de pesos avaliados, mas o índice de robustez é 37,85 e exige cautela.',
    )
    add_body(
        doc,
        'Para a Empresa 2, o perfil de custo mínimo seleciona `empresa2_candidate_005`, com 1 CD, custo de R$ 77.022.539,91 e saving de 9,497399%. Nos perfis equilibrado, conservador e qualidade/serviço, a recomendação muda para `empresa2_candidate_010`, com 2 CDs, custo de R$ 80.636.126,57 e saving de 5,251382%. Esses números são resultados do simulador e não constituem previsão econômica ou validação fiscal.',
    )
    add_note(
        doc,
        'As limitações decisivas são a proxy de transferência da Empresa 1 baseada em dados da Empresa 2, a cobertura incompleta de distâncias, o cálculo de estoque no modo days_wacc_only sem pooling explícito, a armazenagem proporcional quando não há correspondência na tabela e a ausência de calibração histórica dos spreads Monte Carlo.',
    )

    heading(doc, 'Sumário', 1)
    for item in [
        '1 Introdução',
        '2 Problema, objetivos e escopo',
        '3 Fundamentação e decisões de modelagem',
        '4 Metodologia e implementação',
        '5 Resultados e reconciliação',
        '6 Validação, reprodutibilidade e auditoria',
        '7 Conclusão e limitações',
        'Referências',
        'Apêndice A — Matriz de correções',
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
    doc.add_page_break()

    heading(doc, '1 Introdução', 1)
    add_body(
        doc,
        'A configuração da malha logística afeta simultaneamente custo, nível de serviço, concentração operacional, exposição tributária e necessidade de estoque. Em problemas reais, a decisão raramente depende de uma única métrica: o cenário de menor custo pode elevar concentração ou reduzir margem operacional, enquanto uma configuração mais distribuída pode preservar serviço com custo maior.',
    )
    add_body(
        doc,
        'O projeto foi desenvolvido como um simulador web auditável para apoiar a comparação de cenários de duas empresas com bases de dados e premissas próprias. O escopo do Sprint 4 é a entrega da camada de cenários, otimização por enumeração de candidatos, análise de sensibilidade, recomendação e trilha de auditoria; não é uma implementação de otimização global nem uma validação fiscal formal.',
    )
    add_body(
        doc,
        'A revisão apresentada neste documento responde às críticas do relatório e do plano de trabalho recebido. A prioridade foi preservar a estrutura e os dados do projeto, corrigindo apenas inconsistências que afetavam a interpretação, a reprodutibilidade ou a sustentação dos números.',
    )

    heading(doc, '2 Problema, objetivos e escopo', 1)
    heading(doc, '2.1 Problema de decisão', 2)
    add_body(
        doc,
        'O problema consiste em comparar configurações discretas de CDs para cada empresa, observando o custo total e indicadores complementares. Cada empresa deve ser tratada separadamente: a demanda, a matriz de distâncias, as referências de custo e os parâmetros tributários não podem ser misturados sem declaração explícita.',
    )
    heading(doc, '2.2 Objetivos', 2)
    add_body(
        doc,
        'O objetivo geral é fornecer uma comparação reproduzível de cenários de malha logística. Os objetivos específicos são: calcular um baseline rastreável; simular candidatos; decompor custos; aplicar regimes tributários parametrizados; medir qualidade, risco e robustez; expor incerteza por Monte Carlo; reconciliar resultados com referências disponíveis; e registrar as limitações de dados.',
    )
    heading(doc, '2.3 Justificativa para não usar MILP', 2)
    add_body(
        doc,
        'Modelos MILP são adequados quando há dados suficientemente granulares para representar variáveis de decisão, capacidade, atendimento, custos fixos, restrições de fluxo e regras tributárias em uma formulação integrada. No escopo disponível, parte dos custos é proxy, as distâncias não cobrem todos os fluxos e a camada tributária é parametrizada. A enumeração de um espaço discreto de candidatos é, portanto, uma escolha operacional compatível com os dados e com a necessidade de transparência, não uma afirmação de que o problema não possa ser modelado por MILP em uma etapa futura.',
    )

    heading(doc, '3 Fundamentação e decisões de modelagem', 1)
    add_body(
        doc,
        'A modelagem segue a lógica de custo total de malha, em que transporte, armazenagem, distribuição, transferência, estoque e tributos são avaliados de forma conjunta. A literatura de localização de instalações e desenho de redes fornece a base conceitual para comparar configurações, mas a aplicação neste projeto é uma simulação paramétrica apoiada nos dados disponíveis.',
    )
    add_body(
        doc,
        'O núcleo de comparação é determinístico: fixados dados, parâmetros, cenário, regime tributário e seed quando aplicável, o resultado é reproduzível. A camada Monte Carlo altera drivers por amostragem gaussiana limitada e recalcula o cenário. Ela complementa a comparação, mas não transforma as premissas em estimativas empiricamente calibradas.',
    )
    add_note(
        doc,
        'A palavra ‘probabilidade’ neste relatório refere-se à frequência observada nas amostras simuladas. Não deve ser interpretada como probabilidade estatística validada sem histórico adequado.',
    )

    heading(doc, '4 Metodologia e implementação', 1)
    heading(doc, '4.1 Arquitetura', 2)
    add_body(
        doc,
        'O projeto é uma aplicação web estática organizada por fases. A Fase 2 deriva e reconcilia o baseline; a Fase 3 constrói, simula, compara e avalia cenários; a Fase 4 gera candidatos e scores; a Fase 5 consolida stress tests, Monte Carlo, recomendação e audit trail. O motor tributário canônico fica em `assets/js/shared/tax`; os arquivos em `assets/js/core/tax` são aliases de compatibilidade quando necessários.',
    )
    heading(doc, '4.2 Baselines e saving', 2)
    add_body(
        doc,
        'Foram separados o baseline estrutural atual e o baseline comparável. O primeiro preserva o cenário operacional atual e o regime vigente. O segundo preserva o cenário físico, mas usa o mesmo modo e regime tributário dos candidatos. Para um cenário s, o cálculo publicado é saving_abs = baseline_comparável − custo_s e saving_pct = saving_abs / baseline_comparável.',
    )
    heading(doc, '4.3 Monte Carlo', 2)
    add_body(
        doc,
        'A configuração canônica usa 300 iterações, seed 42 e perfil balanced. Os drivers incluem multiplicadores de frete e demanda, dias de estoque, WACC e multiplicador tributário, com choques compartilhados e idiossincráticos, limites e cálculo de percentis. Testes verificam reprodutibilidade com a mesma seed, diferença com seeds diferentes, quantidade de amostras, limites, efeito dos drivers, percentis e probabilidade de saving.',
    )
    heading(doc, '4.4 Transferência da Empresa 1', 2)
    add_body(
        doc,
        'A base da Empresa 1 não contém tarifa própria de transferência suficiente para todos os fluxos. O código mantém uma proxy explícita baseada em médias ponderadas de notas fiscais da Empresa 2, com taxas por UF e fallback de R$ 0,005/kg-km. Essa aproximação afeta os fluxos positivos da Empresa 1 e é registrada no resultado. As distâncias ausentes da matriz não são inventadas; a parcela definida pelo modelo é exposta como fallback de distribuição.',
    )
    add_note(
        doc,
        'A afirmação correta não é ‘não há mistura de premissas’. A afirmação correta é que as empresas são separadas no fluxo de dados e que, na Empresa 1, existe uma proxy cross-company explicitamente identificada, testada e sujeita a sensibilidade.',
    )
    heading(doc, '4.5 Estoque e armazenagem', 2)
    add_body(
        doc,
        'O custo de estoque permanece no modo days_wacc_only: depende de demanda, dias de estoque e WACC. O código não aplica automaticamente 1/√n sobre todo o estoque, pois não há dados suficientes para separar estoque cíclico, estoque de segurança, lead time e variabilidade de demanda. Portanto, o saving não é atribuído à consolidação de estoque. A armazenagem segue a tabela quando há correspondência; na ausência, aplica a aproximação 0,65 + 0,35 × CDs ativos / CDs do baseline, identificada como proxy.',
    )
    heading(doc, '4.6 Scores, risco e robustez', 2)
    add_body(
        doc,
        'Os perfis avaliados são custo mínimo, equilibrado, conservador e qualidade/serviço. Pesos, defaults e thresholds ficam no catálogo de premissas. Risco e robustez são índices de triagem compostos por regras e pesos; não são probabilidades de ocorrência. A normalização min-max é relativa ao conjunto de candidatos avaliado e o vencedor é apenas o melhor cenário dentro desse conjunto.',
    )

    heading(doc, '5 Resultados e reconciliação', 1)
    add_table(
        doc,
        'Tabela 1 — Baselines canônicos por empresa',
        ['Empresa', 'Baseline estrutural atual', 'Baseline comparável reforma 2033'],
        [
            ['Empresa 1', 'R$ 118.694.234,30', 'R$ 156.868.606,74'],
            ['Empresa 2', 'R$ 49.655.885,59', 'R$ 85.105.332,59'],
        ],
        [3.5, 5.2, 6.5],
    )
    add_body(
        doc,
        'A diferença entre os dois baselines decorre do regime tributário, não de uma alteração silenciosa nos custos físicos. Na Empresa 1, o componente físico/logístico do baseline comparável é R$ 37.196.905,49 e o componente tributário é R$ 119.671.701,25. Na Empresa 2, os componentes são R$ 23.574.434,01 e R$ 61.530.898,58, respectivamente.',
    )
    add_table(
        doc,
        'Tabela 2 — Recomendação por perfil',
        ['Empresa', 'Perfil', 'Cenário', 'CDs', 'Custo total', 'Saving', 'Robustez'],
        [
            ['1', 'quatro perfis', 'empresa1_candidate_015', '14', 'R$ 155.296.806,21', '1,001985%', '37,85'],
            ['2', 'custo mínimo', 'empresa2_candidate_005', '1', 'R$ 77.022.539,91', '9,497399%', '99,10'],
            [
                '2',
                'equilibrado, conservador e qualidade/serviço',
                'empresa2_candidate_010',
                '2',
                'R$ 80.636.126,57',
                '5,251382%',
                '82,80',
            ],
        ],
        [2.0, 3.0, 3.2, 1.0, 3.3, 2.0, 1.5],
    )
    add_body(
        doc,
        'A Empresa 1 é sensível à limitação do índice de robustez: o candidato de menor custo relativo não deve ser tratado como recomendação incondicional. Na Empresa 2, a mudança de vencedor demonstra que a decisão depende da função objetivo; por isso, uma única recomendação sem informar o perfil seria incompleta.',
    )
    doc.add_page_break()
    add_table(
        doc,
        'Tabela 3 — Resultado dos quatro perfis de objetivo',
        ['Empresa', 'Perfil', 'Pesos C/S/R/T/E', 'Vencedor', 'Custo', 'Saving', 'Concentração', 'Risco', 'Robustez'],
        [
            [
                '1',
                'custo mínimo',
                '75/5/5/10/5',
                'candidate_015',
                'R$ 155.296.806,21',
                '1,001985%',
                '18,25%',
                'baixo',
                '37,85',
            ],
            [
                '1',
                'equilibrado',
                '30/25/20/15/10',
                'candidate_015',
                'R$ 155.296.806,21',
                '1,001985%',
                '18,25%',
                'baixo',
                '37,85',
            ],
            [
                '1',
                'conservador',
                '15/25/45/5/10',
                'candidate_015',
                'R$ 155.296.806,21',
                '1,001985%',
                '18,25%',
                'baixo',
                '37,85',
            ],
            [
                '1',
                'qualidade/serviço',
                '10/55/20/5/10',
                'candidate_015',
                'R$ 155.296.806,21',
                '1,001985%',
                '18,25%',
                'baixo',
                '37,85',
            ],
            [
                '2',
                'custo mínimo',
                '75/5/5/10/5',
                'candidate_005',
                'R$ 77.022.539,91',
                '9,497399%',
                '53,98%',
                'baixo',
                '99,10',
            ],
            [
                '2',
                'equilibrado',
                '30/25/20/15/10',
                'candidate_010',
                'R$ 80.636.126,57',
                '5,251382%',
                '34,57%',
                'baixo',
                '82,80',
            ],
            [
                '2',
                'conservador',
                '15/25/45/5/10',
                'candidate_010',
                'R$ 80.636.126,57',
                '5,251382%',
                '34,57%',
                'baixo',
                '82,80',
            ],
            [
                '2',
                'qualidade/serviço',
                '10/55/20/5/10',
                'candidate_010',
                'R$ 80.636.126,57',
                '5,251382%',
                '34,57%',
                'baixo',
                '82,80',
            ],
        ],
        [1.2, 2.0, 2.0, 2.0, 2.4, 1.5, 1.6, 1.2, 1.2],
    )
    doc.add_page_break()
    add_table(
        doc,
        'Tabela 4 — Classes de reconciliação',
        ['Classe', 'Erro percentual absoluto', 'Interpretação'],
        [
            ['aligned', 'até 3%', 'alinhado'],
            ['tolerable', '>3% e até 10%', 'tolerável'],
            ['divergent', '>10%', 'divergente'],
            ['benchmark pendente', 'sem referência comparável', 'não calcular score'],
        ],
        [3.5, 4.5, 7.0],
    )
    add_body(
        doc,
        'A reconciliação só calcula erro quando há uma referência comparável e independente. O valor de um workbook de referência não é confundido com a saída simulada. Quando a referência não está disponível para a mesma definição de custo, o status correto é benchmark pendente.',
    )
    add_table(
        doc,
        'Tabela 5 — Reconciliação do baseline',
        ['Empresa', 'Métrica', 'Referência', 'Simulado', 'Erro abs.', 'Erro %', 'Status'],
        [
            ['1', 'operacional/tributário', 'não disponível', 'R$ 118.694.234,30', '—', '—', 'benchmark pendente'],
            ['2', 'transferência', 'R$ 3.148.100,51', 'R$ 3.148.100,51', 'R$ 0,00', '0,000%', 'aligned'],
            ['2', 'distribuição', 'R$ 7.637.802,74', 'R$ 7.637.802,74', 'R$ 0,00', '0,000%', 'aligned'],
            ['2', 'armazenagem', 'R$ 8.647.999,46', 'R$ 8.647.999,46', 'R$ 0,00', '0,000%', 'aligned'],
            ['2', 'estoque', 'R$ 4.140.531,30', 'R$ 4.140.531,30', 'R$ 0,00', '0,000%', 'aligned'],
            ['2', 'tributação', 'R$ 26.081.451,58', 'R$ 26.081.451,58', 'R$ 0,00', '0,000%', 'aligned'],
            ['2', 'frete', 'R$ 10.785.903,25', 'R$ 10.785.903,25', 'R$ 0,00', '0,000%', 'aligned'],
            ['2', 'logística total', 'R$ 23.574.434,01', 'R$ 23.574.434,01', 'R$ 0,00', '0,000%', 'aligned'],
            ['2', 'total com tributos', 'R$ 49.655.885,59', 'R$ 49.655.885,59', 'R$ 0,00', '0,000%', 'aligned'],
        ],
        [1.1, 2.4, 2.8, 2.8, 1.5, 1.3, 2.0],
    )
    add_table(
        doc,
        'Tabela 6 — Decomposição dos cenários canônicos',
        [
            'Empresa/cenário',
            'Transferência',
            'Distribuição',
            'Armazenagem',
            'Estoque',
            'Tributos (impacto parametrizado; proxy E1)',
            'Total',
        ],
        [
            [
                'E1 candidate_015',
                'R$ 5.506.011,19',
                'R$ 13.765.027,97',
                'R$ 11.378.198,43',
                'R$ 4.975.867,38',
                'R$ 119.671.701,25',
                'R$ 155.296.806,21',
            ],
            [
                'E2 candidate_005',
                'R$ 3.148.100,51',
                'R$ 7.637.802,74',
                'R$ 565.206,78',
                'R$ 4.140.531,30',
                'R$ 61.530.898,58',
                'R$ 77.022.539,91',
            ],
            [
                'E2 candidate_010',
                'R$ 3.148.100,51',
                'R$ 7.637.802,74',
                'R$ 4.178.793,44',
                'R$ 4.140.531,30',
                'R$ 61.530.898,58',
                'R$ 80.636.126,57',
            ],
        ],
        [2.3, 2.0, 2.0, 2.0, 1.8, 2.4, 2.5],
    )
    add_table(
        doc,
        'Tabela 7 — Fallbacks e limitações auditáveis',
        ['Elemento', 'Classificação', 'Valor/unidade', 'Condição e efeito'],
        [
            [
                'Transferência E1',
                'proxy',
                'taxas por UF; fallback R$ 0,005/kg-km',
                'distância/tarifa própria ausente; afeta fluxos E1',
            ],
            [
                'Distribuição E1',
                'fallback',
                '40% do custo de distribuição',
                'distância não recuperável; não inventa distância',
            ],
            ['Frete E2', 'fallback', '2,5% da receita', 'linha CIF ausente'],
            ['Armazenagem', 'proxy', '0,65 + 0,35 × CDs/baseline', 'CD sem correspondência na tabela'],
            ['Estoque', 'parâmetro', '45 dias; WACC 15% a.a.', 'modo days_wacc_only; sem pooling'],
            ['Monte Carlo', 'parâmetro', '300; seed 42; spreads manuais', 'camada exploratória'],
            ['Qualidade/risco', 'fallback', 'quality 70; penalidades 18/8', 'dados ausentes'],
            ['Recomendação', 'parâmetro', 'robustez 70/45; MC 65%/50%', 'régua decisória'],
            ['Reconciliação', 'parâmetro', '3%/10%', 'classificação de desvios'],
        ],
        [3.0, 2.0, 4.5, 5.5],
    )
    add_body(
        doc,
        'Na Tabela 6, “Tributos” representa impacto tributário parametrizado, não custo tributário real nem validação fiscal. Na Empresa 1, a parcela também deve ser lida junto da proxy explicitada na Tabela 7. A tabela de resultados original apresentava valores arredondados próximos de R$ 161 milhões, R$ 158 milhões e R$ 159 milhões, mas savings de 1,18% e 1,12% que não eram compatíveis com esses arredondamentos. A causa operacional foi a comparação com uma base não equivalente. Nesta versão, os valores são derivados do baseline comparável e recebem casas decimais suficientes para permitir a conferência.',
    )
    add_table(
        doc,
        'Tabela 8 — Cobertura observada dos fallbacks nos cenários publicados',
        ['Empresa/cenário', 'Fluxos', 'Proxy transferência', 'Distância ausente', 'Fallback receita', 'Estoque'],
        [
            ['E1 candidate_015', '36', '36/36 (100%)', '36/36 (100%)', '0/36', '36 fluxos; sem pooling'],
            ['E2 candidate_005', '76', '0/76', '0/76', '7/76 (9,21%)', '76 fluxos; sem pooling'],
            ['E2 candidate_010', '76', '0/76', '0/76', '0/76', '76 fluxos; sem pooling'],
        ],
        [2.7, 1.0, 2.6, 2.4, 2.2, 3.0],
    )
    add_note(
        doc,
        'Na Empresa 1, a cobertura de distância ausente é total nos fluxos avaliados; por isso, a recomendação é indicativa até que a matriz seja corrigida. Na Empresa 2, o fallback de receita aparece somente nos fluxos em que a tarifa de transferência não está disponível.',
    )
    add_table(
        doc,
        'Tabela 9 — Sensibilidade do saving à proxy de transferência da Empresa 1',
        ['Multiplicador da taxa', 'Saving do candidate_015', 'Interpretação'],
        [
            ['0,75', '1,879473%', 'proxy 25% menor'],
            ['1,00', '1,001985%', 'configuração canônica'],
            ['1,25', '0,124498%', 'proxy 25% maior'],
        ],
        [3.5, 4.5, 5.0],
    )
    heading(doc, '5.1 Resultado Monte Carlo', 2)
    add_body(
        doc,
        'O Monte Carlo deve ser lido como análise exploratória de incerteza. A média, mediana, percentis e probabilidade de saving são calculados sobre amostras limitadas pelos bounds do modelo. Como os spreads não foram calibrados por histórico, o resultado informa sensibilidade às premissas e não uma previsão estatística validada.',
    )
    add_table(
        doc,
        'Tabela 10 — Saída Monte Carlo exploratória da Empresa 2',
        [
            'Configuração',
            'Média do custo',
            'Mediana do custo',
            'P10 saving',
            'P50 saving',
            'P90 saving',
            'P(saving > 0)',
        ],
        [
            [
                '300 iterações; seed 42; balanced',
                'R$ 81.075.432,57',
                'R$ 81.096.647,68',
                '-2,598700%',
                '4,710263%',
                '11,809160%',
                '80,3333%',
            ]
        ],
        [3.3, 2.4, 2.4, 1.7, 1.7, 1.7, 2.0],
    )
    add_note(
        doc,
        'Os números desta tabela são exploratórios e não fundamentam os totais determinísticos nem a recomendação principal. A execução da Empresa 1 também é registrada nos artefatos de prova, mas não é publicada aqui como previsão; os spreads continuam sem calibração histórica.',
    )
    heading(doc, '5.2 Sensibilidade da recomendação', 2)
    add_body(
        doc,
        'Na Empresa 1, a sensibilidade da taxa de transferência da proxy, com multiplicadores de 0,75, 1,00 e 1,25, produz savings de aproximadamente 1,879473%, 1,001985% e 0,124498% para o candidato canônico. Essa faixa evidencia que a conclusão econômica depende da premissa cross-company e deve ser confirmada com dados próprios antes de uma decisão de implantação.',
    )

    heading(doc, '6 Validação, reprodutibilidade e auditoria', 1)
    add_body(
        doc,
        'A validação inclui instalação limpa, padrões de código, contratos de módulos, paths, criptografia, servidor HTTP, lógica das fases, tributação, Monte Carlo, score, estoque, fallbacks, reconciliação, exportação, E2E desktop/mobile, rotas inválidas, localStorage corrompido e bloqueio por senha. Os testes verificam valores e relações matemáticas, não apenas a existência de arquivos.',
    )
    add_body(
        doc,
        'A branch de preservação mantém o estado anterior à correção. A branch de trabalho concentra commits semânticos e não contém a senha dos dados. O CI executa a suíte pública automaticamente e falha explicitamente quando o secret protegido não está configurado; a suíte protegida só pode ser aprovada em ambiente com credencial legítima.',
    )
    add_note(
        doc,
        'As Rodadas 1 e 2 só são consideradas aprovadas quando executadas após o último commit, em ambientes novos, sem falhas, sem warnings sem justificativa e com inspeção visual de todas as páginas do PDF revisado.',
    )

    heading(doc, '7 Conclusão e limitações', 1)
    add_body(
        doc,
        'A versão revisada entrega uma comparação de cenários mais consistente porque alinha baseline e regime tributário, explicita proxies, remove a atribuição indevida de pooling de estoque, centraliza fallbacks e verifica a aritmética dos savings. A decisão para a Empresa 1 deve ser tratada como indicativa e condicionada à substituição da proxy de transferência e à melhoria da matriz de distâncias. Para a Empresa 2, a configuração recomendada depende do perfil de decisão, com vantagem de custo para 1 CD e maior equilíbrio para 2 CDs.',
    )
    add_body(
        doc,
        'Permanecem limitações legítimas: ausência de calibração histórica do Monte Carlo; qualidade e robustez como índices parametrizados; cobertura incompleta de distâncias; armazenagem proporcional em correspondências ausentes; estoque sem pooling explícito; e camada tributária como impacto parametrizado, não validação fiscal. Essas limitações são parte do resultado e devem acompanhar qualquer uso do simulador.',
    )

    heading(doc, 'Referências', 1)
    refs = [
        'BALLOU, Ronald H. Gerenciamento da cadeia de suprimentos/logística empresarial. Porto Alegre: Bookman, 2006.',
        'BANKS, Jerry et al. Discrete-event system simulation. 5. ed. Upper Saddle River: Prentice Hall, 2010.',
        'BRASIL. Emenda Constitucional n. 132, de 20 de dezembro de 2023. Brasília, DF, 2023.',
        'CHOPRA, Sunil; MEINDL, Peter. Supply chain management: strategy, planning, and operation. 6. ed. Boston: Pearson, 2016.',
        'EPPEN, Gary D. Effects of centralization on expected costs in a multi-location newsboy problem. Management Science, v. 25, n. 5, p. 498–501, 1979.',
        'GEOFFRION, Arthur M.; GRAVES, George W. Multicommodity distribution system design by Benders decomposition. Management Science, v. 20, n. 5, p. 822–844, 1974.',
        'KLOSE, Andreas; DREXL, Andreas. Facility location models for distribution system design. European Journal of Operational Research, v. 162, n. 1, p. 4–29, 2005.',
        'LAW, Averill M. Simulation modeling and analysis. 5. ed. New York: McGraw-Hill, 2015.',
        'MELO, M. T.; NICKEL, S.; SALDANHA-DA-GAMA, F. Facility location and supply chain management: a review. European Journal of Operational Research, v. 196, n. 2, p. 401–412, 2009.',
        'RECEITA FEDERAL. Reforma tributária: materiais institucionais e legislação aplicável. Brasília, 2024.',
        'SALTELLI, Andrea et al. Global sensitivity analysis: the primer. Chichester: Wiley, 2008.',
        'SCHOEMAKER, Paul J. H. Scenario planning: a tool for strategic thinking. Sloan Management Review, v. 36, n. 2, p. 25–40, 1995.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()
    heading(doc, 'Apêndice A — Matriz de correções', 1)
    add_table(
        doc,
        'Tabela 11 — Rastreabilidade da revisão',
        ['Crítica', 'Arquivo/módulo', 'Resposta', 'Regressão'],
        [
            [
                'Monte Carlo',
                'monte-carlo-engine.js',
                'camada probabilística complementar, drivers e seed auditáveis',
                'test_phase3_logic',
            ],
            [
                'Calibração cruzada',
                'physical-cost-engine.js',
                'proxy E1 explícita e sensibilidade',
                'test_phase3_assumptions_logic',
            ],
            [
                'Estoque',
                'physical-cost-engine.js',
                'days_wacc_only e warning de pooling',
                'test_phase3_assumptions_logic',
            ],
            [
                'Tributação',
                'shared/tax',
                'fonte canônica; três arquivos mortos removidos; aliases necessários',
                'test_tax_aliases_and_regimes',
            ],
            [
                'Baseline',
                'optimizer-utils.js; reconciliation-engine.js',
                'mesmo regime e régua única',
                'test_phase4_scoring_logic; test_phase2_reconciliation_logic',
            ],
            ['Entrega', 'README; quality.yml', 'execução pública/protegida e CI', 'check_package; code_standards'],
        ],
        [3.0, 4.0, 6.0, 3.0],
    )

    out = OUT / 'R_SPRINT4_Grupo2_Visagio_revisado.docx'
    doc.save(out)
    print(out)


if __name__ == '__main__':
    main()
